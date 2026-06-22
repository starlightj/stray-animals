"""
生成「软件项目管理 任务一」Word 文档
任务一：网站设计构建及整体验收测试文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ====== 封面 ======
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('软件项目管理')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('课程设计任务一')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0xD3, 0x54, 0x00)

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('网站设计构建及整体验收测试文档')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— 基于图像识别的校园流浪动物管理系统')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x63, 0x6E, 0x72)

for _ in range(6):
    doc.add_paragraph('')

info = [('课程名称','软件项目管理'),('项目名称','基于图像识别的校园流浪动物管理系统'),
        ('小组成员','__________、__________、__________'),('指导教师','__________'),('提交日期','2026年6月')]
table = doc.add_table(rows=len(info), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
for i,(k,v) in enumerate(info):
    for j,t in enumerate([k,v]):
        c = table.cell(i,j); c.text = t
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ====== 正文 ======
doc.add_heading('任务一：网站设计构建及整体验收测试文档', level=1)
doc.add_paragraph('本任务要求完成校园流浪动物管理系统的网站设计构建、环境部署，并对网站功能进行全面测试验收。以下从网站功能介绍、环境搭建部署过程、测试的必要性与流程、测试对象与需求、测试目标等几个方面进行阐述。')

# 一、网站功能介绍
doc.add_heading('一、网站功能介绍', level=1)
doc.add_paragraph('本系统是一个基于图像识别的校园流浪动物管理平台，共包含七个核心功能模块：')

funcs = [
 ('📊 数据看板','系统首页，展示总动物数量、物种数量、识别次数等统计指标，以及最新登记的动物卡片，提供快速入口导航。'),
 ('📷 AI图像识别','用户上传动物照片，前端加载TensorFlow.js+MobileNet模型实时推理，自动识别猫/狗并显示置信度。低于60%提示"不确定，请手动选择"。结果自动填充保存表单。'),
 ('📋 动物档案','卡片网格展示所有动物信息，支持按名称/特征搜索和按物种筛选，点击查看完整详情。'),
 ('🗺️ 地图追踪','列表展示所有动物的发现位置和地址信息，支持按物种筛选查看分布。'),
 ('🏠 动物领养','展示可供领养的动物，用户可提交领养申请（姓名、联系方式、理由），申请记录存入数据库。'),
 ('💬 评论系统','选择动物查看相关评论，支持发表评论（昵称选填，内容必填），按时间倒序排列。'),
 ('⚙️ 管理后台','支持动物增删改查，表格展示所有记录，支持搜索筛选、编辑弹窗（含图片上传）、删除确认。'),
]
for t,d in funcs:
    p = doc.add_paragraph()
    p.add_run(t+'：').font.bold = True
    p.add_run(d)
    p.paragraph_format.space_after = Pt(3)

# 二、环境搭建
doc.add_heading('二、环境搭建与部署过程', level=1)

doc.add_heading('2.1 开发环境', level=2)
for k,v in [('操作系统','Windows 11'),('运行环境','Node.js 24.x'),('开发工具','VS Code + 微信开发者工具'),('版本控制','Git + GitHub')]:
    p = doc.add_paragraph(); p.add_run(k+'：').font.bold = True; p.add_run(v)

doc.add_heading('2.2 部署步骤', level=2)
for i,s in enumerate([
 '克隆代码：git clone https://github.com/starlightj/stray-animals.git',
 '进入目录：cd stray-animals',
 '安装后端依赖：cd backend && npm install',
 '启动服务：node server-debug.js（自动创建SQLite数据库并初始化数据）',
 '访问系统：浏览器打开 http://localhost:3000',
],1):
    doc.add_paragraph(f'步骤{i}：{s}')

doc.add_heading('2.3 技术栈', level=2)
for k,v in [('前端','HTML5+CSS3+JavaScript+TensorFlow.js+MobileNet'),
            ('后端','Node.js+Express+Multer'),
            ('数据库','SQLite(better-sqlite3，零配置)'),
            ('AI识别','MobileNet预训练模型，前端直接推理')]:
    p = doc.add_paragraph(); p.add_run(k+'：').font.bold = True; p.add_run(v)

# 三、测试必要性与流程
doc.add_heading('三、测试的必要性与流程', level=1)

doc.add_heading('3.1 测试必要性', level=2)
doc.add_paragraph('软件测试是保证系统质量的关键环节，其必要性体现在：')
for r in ['保证功能完整性：验证所有功能模块是否按需求正确实现',
          '发现潜在缺陷：发现代码中的逻辑错误、边界问题和异常处理缺陷',
          '验证用户体验：检查界面是否友好、交互是否流畅、反馈是否及时',
          '确保数据安全：验证数据存储的正确性、一致性和完整性',
          '降低维护成本：尽早发现问题可大幅降低后期修复成本']:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('3.2 测试流程', level=2)
for t,d in [('制定测试计划','确定测试范围、目标、资源和时间安排'),
            ('设计测试用例','针对每个模块编写详细测试用例，含正常和异常流程'),
            ('执行功能测试','按测试用例逐项执行，记录测试结果和问题'),
            ('记录缺陷','对发现的问题进行记录、分类和优先级评定'),
            ('回归测试','修复问题后重新测试，确保修复有效'),
            ('编写测试报告','汇总测试结果，形成验收测试文档')]:
    p = doc.add_paragraph(); p.add_run(t+'：').font.bold = True; p.add_run(d)

# 四、测试对象与需求
doc.add_heading('四、测试的对象及测试需求', level=1)

doc.add_heading('4.1 测试对象', level=2)
doc.add_paragraph('本系统的测试对象包括以下七个功能模块：')
for m in ['数据看板模块','AI图像识别模块','动物档案模块','地图追踪模块','动物领养模块','评论系统模块','管理后台模块']:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('4.2 测试环境需求', level=2)
for k,v in [('硬件','普通PC，CPU i5及以上，内存8GB及以上'),('软件','Windows 10/11，Chrome/Firefox/Edge浏览器'),('网络','局域网或本地回环（localhost）')]:
    p = doc.add_paragraph(); p.add_run(k+'需求：').font.bold = True; p.add_run(v)

doc.add_heading('4.3 测试数据需求', level=2)
doc.add_paragraph('系统初始化时自动插入4条动物示例数据（大黄/狗、小黑/猫、小花/猫、旺财/狗），以及对应的评论示例数据（8条）和领养申请示例数据（2条）。测试过程中可在此基础上添加、修改和删除数据以验证各功能。')

# 五、测试目标
doc.add_heading('五、测试的需求及目标', level=1)

doc.add_heading('5.1 功能测试目标', level=2)
for k,v in [('完整性','所有功能模块均可正常使用，无缺失功能'),
            ('正确性','各功能执行结果与预期一致，数据处理准确'),
            ('健壮性','对异常输入和非法操作有正确处理和提示'),
            ('一致性','不同模块间的数据保持一致，无冲突'),
            ('易用性','界面简洁明了，操作直观，反馈及时')]:
    p = doc.add_paragraph(); p.add_run(k+'：').font.bold = True; p.add_run(v)

doc.add_heading('5.2 验收标准', level=2)
doc.add_paragraph('所有基本功能测试通过，无严重Bug（影响核心功能正常使用的缺陷），界面无明显错乱，数据读写正确，即可通过验收。')

# 六、验收测试表
doc.add_heading('六、基本功能测试验收表', level=1)
doc.add_paragraph('')

items = [
 ('1','源代码与文档交付','项目源代码完整可编译','GitHub仓库包含全部代码'),
 ('2','源代码与文档交付','文档资料齐全','任务文档、答辩PPT等齐全'),
 ('3','安装部署与调试','环境搭建步骤清晰','按README可完成部署'),
 ('4','安装部署与调试','服务端正常启动','node server-debug.js启动成功'),
 ('5','安装部署与调试','数据库自动初始化','首次启动自动建表和插入数据'),
 ('6','与需求一致程度','功能覆盖全部需求','7个模块全部实现'),
 ('7','与需求一致程度','AI识别功能可用','上传图片可识别猫/狗'),
 ('8','性能表现','页面加载速度','3秒内完成首屏加载'),
 ('9','性能表现','图片上传响应','5秒内完成上传并返回'),
 ('10','安全性','SQL注入防护','参数化查询防止注入'),
 ('11','安全性','文件上传安全','限制文件类型和大小'),
 ('12','界面简洁与易用性','UI风格统一','暖色调动物主题风格一致'),
 ('13','界面简洁与易用性','导航清晰','7个导航入口明确'),
 ('14','界面简洁与易用性','响应式适配','手机端和PC端均可正常显示'),
 ('15','详细功能列表','数据看板','统计数据和最新动物展示'),
 ('16','详细功能列表','AI图像识别上报','上传→识别→保存全流程'),
 ('17','详细功能列表','动物档案','搜索/筛选/查看详情'),
 ('18','详细功能列表','地图追踪','按物种筛选查看位置'),
 ('19','详细功能列表','动物领养','提交领养申请'),
 ('20','详细功能列表','评论系统','发表和查看评论'),
 ('21','详细功能列表','管理后台','增删改查+图片上传'),
]

t2 = doc.add_table(rows=len(items)+2, cols=5)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

# 合并表头
for ci in range(5):
    t2.cell(0,ci).merge(t2.cell(1,ci))

for i,h in enumerate(['序号','内容模块','具体内容','说明','验收结果']):
    c = t2.cell(0,i); c.text = h
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.paragraphs[0].runs: r.font.bold = True; r.font.size = Pt(10)

for ri,(s,m,c,d) in enumerate(items):
    for ci,v in enumerate([s,m,c,d,'□ 是  □ 否']):
        cell = t2.cell(ri+2,ci); cell.text = v
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cell.paragraphs[0].runs: r.font.size = Pt(9.5)

out = r'c:\Users\明卿\Desktop\stray-animals\任务一_需求分析文档.docx'
doc.save(out)
print(f'✅ 已生成: {out}')
